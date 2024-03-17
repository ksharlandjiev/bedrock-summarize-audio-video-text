from docx import Document
from handlers.abstract_handler import AbstractHandler

class MicrosoftWordReaderHandler(AbstractHandler):
    
    def handle(self, request: dict) -> dict:
        print("Processing DOCX file...")

        # Initialize a variable to hold the extracted text
        text_content = ''

        # Load the document from the path specified in the request
        doc = Document(request.get("path", None))

        # Iterate through each paragraph in the document to extract text
        for para in doc.paragraphs:
            text_content += para.text + '\n'
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Iterate through paragraphs inside each cell
                    for paragraph in cell.paragraphs:
                        text_content += paragraph.text + '\n'

        # Update the request with the extracted text
        request.update({"text": text_content})

        # Call the next handler in the chain
        return super().handle(request)
